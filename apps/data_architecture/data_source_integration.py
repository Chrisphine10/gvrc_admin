"""
Extensible Data Source Integration Framework
Supports multiple data sources with pluggable adapters
"""

import logging
import json
import csv
import requests
import hashlib
from typing import Dict, List, Any, Optional, Callable
from abc import ABC, abstractmethod
from datetime import datetime
from django.conf import settings
from .models import DataSource, RawDataRecord, DataProcessingEvent
from .raw_data_lake import RawDataLake
from .data_validation import DataValidator
from .ai_geolocation import AIGeolocationEnhancer

logger = logging.getLogger(__name__)


class DataSourceAdapter(ABC):
    """Abstract base class for data source adapters"""
    
    def __init__(self, source_config: Dict[str, Any]):
        self.source_config = source_config
        self.validator = DataValidator()
        self.geolocation_enhancer = AIGeolocationEnhancer()
    
    @abstractmethod
    def connect(self) -> bool:
        """Establish connection to data source"""
        pass
    
    @abstractmethod
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from source"""
        pass
    
    @abstractmethod
    def get_schema(self) -> Dict[str, Any]:
        """Get data schema information"""
        pass
    
    def validate_data(self, data: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Validate extracted data"""
        # Determine data type based on content
        data_type = 'facility'  # default
        if data and 'organization_name' in data[0]:
            data_type = 'gbv_organization'
        elif data and 'facility_name' in data[0]:
            data_type = 'facility'
        
        return self.validator.validate_batch(data, data_type)
    
    def enhance_geolocation(self, data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Enhance data with geolocation information"""
        enhanced_data = []
        for record in data:
            try:
                # Extract address components for geolocation
                address = record.get('facility_name', '') or record.get('text', '')
                county = record.get('county', '')
                constituency = record.get('constituency', '')
                ward = record.get('ward', '')
                
                # Enhance coordinates
                geo_result = self.geolocation_enhancer.enhance_coordinates(
                    address, county, constituency, ward
                )
                
                if geo_result:
                    record.update(geo_result)
                
                enhanced_data.append(record)
            except Exception as e:
                logger.warning(f"Geolocation enhancement failed for record: {str(e)}")
                enhanced_data.append(record)
        return enhanced_data


class CSVDataSourceAdapter(DataSourceAdapter):
    """CSV file data source adapter"""
    
    def connect(self) -> bool:
        """Check if CSV file exists and is accessible"""
        try:
            file_path = self.source_config.get('file_path')
            if not file_path:
                return False
            
            with open(file_path, 'r') as f:
                # Try to read first line to verify file is readable
                f.readline()
            return True
        except Exception as e:
            logger.error(f"CSV connection failed: {str(e)}")
            return False
    
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from CSV file"""
        try:
            file_path = self.source_config.get('file_path')
            encoding = self.source_config.get('encoding', 'utf-8')
            delimiter = self.source_config.get('delimiter', ',')
            
            data = []
            with open(file_path, 'r', encoding=encoding) as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                for i, row in enumerate(reader):
                    if limit and i >= limit:
                        break
                    
                    # Clean and prepare data
                    cleaned_row = {k.strip(): v.strip() if isinstance(v, str) else v 
                                 for k, v in row.items()}
                    data.append(cleaned_row)
            
            logger.info(f"Extracted {len(data)} records from CSV: {file_path}")
            return data
            
        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}")
            return []
    
    def get_schema(self) -> Dict[str, Any]:
        """Get CSV schema information"""
        try:
            file_path = self.source_config.get('file_path')
            with open(file_path, 'r') as f:
                reader = csv.DictReader(f)
                columns = reader.fieldnames or []
            
            return {
                'type': 'csv',
                'columns': columns,
                'file_path': file_path,
                'total_columns': len(columns)
            }
        except Exception as e:
            logger.error(f"Schema extraction failed: {str(e)}")
            return {'type': 'csv', 'columns': [], 'error': str(e)}


class JSONDataSourceAdapter(DataSourceAdapter):
    """JSON file data source adapter"""
    
    def connect(self) -> bool:
        """Check if JSON file exists and is valid"""
        try:
            file_path = self.source_config.get('file_path')
            if not file_path:
                return False
            
            with open(file_path, 'r') as f:
                json.load(f)
            return True
        except Exception as e:
            logger.error(f"JSON connection failed: {str(e)}")
            return False
    
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from JSON file"""
        try:
            file_path = self.source_config.get('file_path')
            array_key = self.source_config.get('array_key', None)
            
            with open(file_path, 'r') as f:
                data = json.load(f)
            
            # Handle different JSON structures
            if array_key and array_key in data:
                records = data[array_key]
            elif isinstance(data, list):
                records = data
            else:
                records = [data]  # Single object
            
            # Apply limit if specified
            if limit:
                records = records[:limit]
            
            logger.info(f"Extracted {len(records)} records from JSON: {file_path}")
            return records
            
        except Exception as e:
            logger.error(f"JSON extraction failed: {str(e)}")
            return []
    
    def get_schema(self) -> Dict[str, Any]:
        """Get JSON schema information"""
        try:
            file_path = self.source_config.get('file_path')
            with open(file_path, 'r') as f:
                data = json.load(f)
            
            # Analyze structure
            if isinstance(data, list) and data:
                sample_record = data[0]
                columns = list(sample_record.keys()) if isinstance(sample_record, dict) else []
            elif isinstance(data, dict):
                columns = list(data.keys())
            else:
                columns = []
            
            return {
                'type': 'json',
                'columns': columns,
                'file_path': file_path,
                'total_columns': len(columns)
            }
        except Exception as e:
            logger.error(f"Schema extraction failed: {str(e)}")
            return {'type': 'json', 'columns': [], 'error': str(e)}


class APIDataSourceAdapter(DataSourceAdapter):
    """REST API data source adapter"""
    
    def connect(self) -> bool:
        """Test API connection"""
        try:
            url = self.source_config.get('base_url')
            if not url:
                return False
            
            # Test with a simple request
            response = requests.get(url, timeout=10)
            return response.status_code == 200
        except Exception as e:
            logger.error(f"API connection failed: {str(e)}")
            return False
    
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from API"""
        try:
            base_url = self.source_config.get('base_url')
            endpoint = self.source_config.get('endpoint', '')
            headers = self.source_config.get('headers', {})
            params = self.source_config.get('params', {})
            
            # Apply limit as parameter if supported
            if limit:
                params['limit'] = limit
            
            url = f"{base_url.rstrip('/')}/{endpoint.lstrip('/')}"
            response = requests.get(url, headers=headers, params=params, timeout=30)
            response.raise_for_status()
            
            data = response.json()
            
            # Handle different response structures
            if isinstance(data, dict):
                # Look for common data keys
                records = data.get('data', data.get('results', data.get('items', [data])))
            else:
                records = data
            
            if not isinstance(records, list):
                records = [records]
            
            logger.info(f"Extracted {len(records)} records from API: {url}")
            return records
            
        except Exception as e:
            logger.error(f"API extraction failed: {str(e)}")
            return []
    
    def get_schema(self) -> Dict[str, Any]:
        """Get API schema information"""
        try:
            base_url = self.source_config.get('base_url')
            schema_endpoint = self.source_config.get('schema_endpoint', 'schema')
            
            url = f"{base_url.rstrip('/')}/{schema_endpoint.lstrip('/')}"
            response = requests.get(url, timeout=10)
            
            if response.status_code == 200:
                schema = response.json()
                return {
                    'type': 'api',
                    'base_url': base_url,
                    'schema': schema
                }
            else:
                return {
                    'type': 'api',
                    'base_url': base_url,
                    'error': f"Schema endpoint returned {response.status_code}"
                }
        except Exception as e:
            logger.error(f"Schema extraction failed: {str(e)}")
            return {'type': 'api', 'error': str(e)}


class DatabaseDataSourceAdapter(DataSourceAdapter):
    """Database data source adapter"""
    
    def connect(self) -> bool:
        """Test database connection"""
        try:
            from django.db import connections
            connection_name = self.source_config.get('connection_name', 'default')
            connection = connections[connection_name]
            connection.ensure_connection()
            return True
        except Exception as e:
            logger.error(f"Database connection failed: {str(e)}")
            return False
    
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from database"""
        try:
            from django.db import connections
            connection_name = self.source_config.get('connection_name', 'default')
            query = self.source_config.get('query')
            
            if not query:
                logger.error("No query specified for database adapter")
                return []
            
            # Add limit to query if specified
            if limit:
                query = f"{query} LIMIT {limit}"
            
            connection = connections[connection_name]
            with connection.cursor() as cursor:
                cursor.execute(query)
                columns = [col[0] for col in cursor.description]
                rows = cursor.fetchall()
                
                data = [dict(zip(columns, row)) for row in rows]
            
            logger.info(f"Extracted {len(data)} records from database")
            return data
            
        except Exception as e:
            logger.error(f"Database extraction failed: {str(e)}")
            return []
    
    def get_schema(self) -> Dict[str, Any]:
        """Get database schema information"""
        try:
            from django.db import connections
            connection_name = self.source_config.get('connection_name', 'default')
            table_name = self.source_config.get('table_name')
            
            if not table_name:
                return {'type': 'database', 'error': 'No table name specified'}
            
            connection = connections[connection_name]
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT column_name, data_type, is_nullable
                    FROM information_schema.columns
                    WHERE table_name = %s
                    ORDER BY ordinal_position
                """, [table_name])
                
                columns = cursor.fetchall()
            
            return {
                'type': 'database',
                'table_name': table_name,
                'columns': [{'name': col[0], 'type': col[1], 'nullable': col[2]} for col in columns]
            }
        except Exception as e:
            logger.error(f"Schema extraction failed: {str(e)}")
            return {'type': 'database', 'error': str(e)}


class DataSourceIntegrationManager:
    """Manages multiple data source integrations"""
    
    def __init__(self):
        self.adapters = {
            'csv': CSVDataSourceAdapter,
            'json': JSONDataSourceAdapter,
            'api': APIDataSourceAdapter,
            'database': DatabaseDataSourceAdapter,
            'pdf': PDFDataSourceAdapter,
            'docx': DOCXDataSourceAdapter,
            'text': TextDataSourceAdapter,
        }
        self.raw_data_lake = RawDataLake()
    
    def register_adapter(self, source_type: str, adapter_class: type):
        """Register a new data source adapter"""
        self.adapters[source_type] = adapter_class
        logger.info(f"Registered adapter for source type: {source_type}")
    
    def create_data_source(self, name: str, source_type: str, config: Dict[str, Any]) -> DataSource:
        """Create and register a new data source"""
        try:
            data_source = DataSource.objects.create(
                name=name,
                source_type=source_type,
                configuration=config,
                is_active=True
            )
            logger.info(f"Created data source: {name} ({source_type})")
            return data_source
        except Exception as e:
            logger.error(f"Failed to create data source: {str(e)}")
            raise
    
    def ingest_from_source(self, data_source: DataSource, limit: Optional[int] = None) -> Dict[str, Any]:
        """Ingest data from a registered data source"""
        try:
            # Get adapter for source type
            adapter_class = self.adapters.get(data_source.source_type)
            if not adapter_class:
                raise ValueError(f"No adapter found for source type: {data_source.source_type}")
            
            # Create adapter instance
            adapter = adapter_class(data_source.configuration)
            
            # Test connection
            if not adapter.connect():
                raise ConnectionError(f"Failed to connect to data source: {data_source.name}")
            
            # Extract data
            raw_data = adapter.extract_data(limit)
            if not raw_data:
                logger.warning(f"No data extracted from source: {data_source.name}")
                return {'status': 'warning', 'message': 'No data extracted', 'records': 0}
            
            # Enhance with geolocation first
            enhanced_data = adapter.enhance_geolocation(raw_data)
            
            # Validate enhanced data
            validation_results = adapter.validate_data(enhanced_data)
            
            # Store in raw data lake
            stored_records = []
            for record in enhanced_data:
                # Generate unique data ID and checksum
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
                data_id = f"{data_source.id}_{len(stored_records)}_{timestamp}"
                # Include timestamp in checksum to ensure uniqueness
                checksum_data = json.dumps(record, sort_keys=True) + timestamp
                checksum = hashlib.md5(checksum_data.encode()).hexdigest()
                
                raw_record = RawDataRecord.objects.create(
                    source=data_source,
                    data_id=data_id,
                    raw_data=record,
                    checksum=checksum,
                    file_path=data_source.configuration.get('file_path', ''),
                    processing_status='ingested'
                )
                stored_records.append(raw_record.id)
            
            # Log processing event
            DataProcessingEvent.objects.create(
                event_type='data_ingestion',
                record_id=f"batch_{data_source.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                source=data_source,
                event_data={
                    'records_processed': len(stored_records),
                    'validation_results': validation_results,
                    'processing_time': datetime.now().isoformat()
                },
                success=True
            )
            
            logger.info(f"Successfully ingested {len(stored_records)} records from {data_source.name}")
            return {
                'status': 'success',
                'message': f'Ingested {len(stored_records)} records',
                'records': len(stored_records),
                'validation_results': validation_results
            }
            
        except Exception as e:
            logger.error(f"Data ingestion failed for {data_source.name}: {str(e)}")
            
            # Log error event
            DataProcessingEvent.objects.create(
                event_type='data_ingestion',
                record_id=f"error_{data_source.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                source=data_source,
                event_data={'error': str(e)},
                success=False,
                error_message=str(e)
            )
            
            return {
                'status': 'error',
                'message': str(e),
                'records': 0
            }
    
    def get_available_sources(self) -> List[Dict[str, Any]]:
        """Get list of available data sources"""
        sources = DataSource.objects.filter(is_active=True)
        return [
            {
                'id': source.id,
                'name': source.name,
                'type': source.source_type,
                'created_at': source.created_at.isoformat(),
                'last_sync': source.last_sync.isoformat() if source.last_sync else None
            }
            for source in sources
        ]
    
    def test_source_connection(self, data_source: DataSource) -> Dict[str, Any]:
        """Test connection to a data source"""
        try:
            adapter_class = self.adapters.get(data_source.source_type)
            if not adapter_class:
                return {'status': 'error', 'message': f'No adapter for type: {data_source.source_type}'}
            
            adapter = adapter_class(data_source.configuration)
            connected = adapter.connect()
            
            if connected:
                schema = adapter.get_schema()
                return {
                    'status': 'success',
                    'message': 'Connection successful',
                    'schema': schema
                }
            else:
                return {
                    'status': 'error',
                    'message': 'Connection failed'
                }
        except Exception as e:
            return {
                'status': 'error',
                'message': str(e)
            }


class PDFDataSourceAdapter(DataSourceAdapter):
    """PDF file data source adapter"""
    
    def connect(self) -> bool:
        """Check if PDF file exists and is readable"""
        try:
            file_path = self.source_config.get('file_path')
            if not file_path:
                return False
            
            import PyPDF2
            with open(file_path, 'rb') as f:
                PyPDF2.PdfReader(f)
            return True
        except Exception as e:
            logger.error(f"PDF connection failed: {str(e)}")
            return False
    
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from PDF file"""
        try:
            import PyPDF2
            import re
            import os
            
            file_path = self.source_config.get('file_path')
            data = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    if limit and len(data) >= limit:
                        break
                        
                    text = page.extract_text()
                    
                    # For KMPDC facilities PDF, look for facility patterns
                    if 'facility' in file_path.lower() or 'kmdpc' in file_path.lower() or 'licensed' in file_path.lower():
                        facilities = self._extract_facilities_from_pdf_text(text)
                        data.extend(facilities)
                    else:
                        # General text extraction
                        lines = text.split('\n')
                        for line in lines:
                            if line.strip() and len(line.split()) > 3:
                                data.append({
                                    'page': page_num + 1,
                                    'text': line.strip(),
                                    'source_file': os.path.basename(file_path)
                                })
            
            return data[:limit] if limit else data
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            return []
    
    def _extract_facilities_from_pdf_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract facility information from PDF text using regex patterns"""
        import re
        facilities = []
        
        # Split text into lines and process
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 10:  # Skip very short lines
                continue
            
            # Look for lines that contain facility information
            # KMPDC format: S/NO Reg_No Facility_Name Facility_Type Facility_Agent Level County Status
            if re.match(r'^\d+[A-Z]+-\d+', line):  # Lines starting with registration number
                parts = line.split()
                if len(parts) >= 7:  # Should have at least 7 parts
                    try:
                        facility_data = {
                            's_no': parts[0],
                            'registration_number': parts[1],
                            'facility_name': ' '.join(parts[2:-5]),  # Everything between reg_no and last 5 parts
                            'facility_type': parts[-5],
                            'facility_agent': parts[-4],
                            'level': parts[-3],
                            'county': parts[-2],
                            'status': parts[-1],
                            'source_type': 'pdf_extraction'
                        }
                        facilities.append(facility_data)
                    except (IndexError, ValueError):
                        # If parsing fails, try a simpler approach
                        facility_data = {
                            'raw_text': line,
                            'source_type': 'pdf_extraction'
                        }
                        facilities.append(facility_data)
            
            # Also look for facility names that don't follow the exact format
            elif re.search(r'(?:Hospital|Clinic|Health Center|Dispensary|Medical Center|Health Centre)', line, re.IGNORECASE):
                facility_data = {
                    'facility_name': line,
                    'source_type': 'pdf_extraction'
                }
                facilities.append(facility_data)
        
        return facilities
    
    def get_schema(self) -> Dict[str, Any]:
        """Get PDF schema information"""
        try:
            file_path = self.source_config.get('file_path')
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Extract text from first page to analyze structure
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text()
                
                return {
                    'type': 'pdf',
                    'pages': num_pages,
                    'sample_text': text[:500],
                    'file_size_mb': round(os.path.getsize(file_path) / (1024 * 1024), 2)
                }
        except Exception as e:
            return {'type': 'pdf', 'error': str(e)}


class DOCXDataSourceAdapter(DataSourceAdapter):
    """DOCX file data source adapter"""
    
    def connect(self) -> bool:
        """Check if DOCX file exists and is readable"""
        try:
            file_path = self.source_config.get('file_path')
            if not file_path:
                return False
            
            from docx import Document
            Document(file_path)
            return True
        except Exception as e:
            logger.error(f"DOCX connection failed: {str(e)}")
            return False
    
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from DOCX file"""
        try:
            from docx import Document
            import re
            import os
            
            file_path = self.source_config.get('file_path')
            data = []
            doc = Document(file_path)
            
            # Extract text from paragraphs with better filtering
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) > 5:  # Only meaningful paragraphs
                    # Look for organization patterns
                    if re.search(r'(?:FIDA|Organization|Support|Legal|Psychological|Child Protection|GBV)', text, re.IGNORECASE):
                        data.append({
                            'text': text,
                            'type': 'organization_info',
                            'source_file': os.path.basename(file_path)
                        })
                    else:
                        data.append({
                            'text': text,
                            'type': 'paragraph',
                            'source_file': os.path.basename(file_path)
                        })
            
            # Extract data from tables
            for table_num, table in enumerate(doc.tables):
                if table.rows:
                    # Get headers from first row
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    
                    # Extract data from remaining rows
                    for row in table.rows[1:]:
                        row_data = {}
                        for i, cell in enumerate(row.cells):
                            if i < len(headers):
                                row_data[headers[i]] = cell.text.strip()
                        if row_data:
                            row_data['table_number'] = table_num + 1
                            row_data['source_file'] = os.path.basename(file_path)
                            data.append(row_data)
            
            return data[:limit] if limit else data
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return []
    
    def get_schema(self) -> Dict[str, Any]:
        """Get DOCX schema information"""
        try:
            file_path = self.source_config.get('file_path')
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text and analyze structure
            full_text = []
            tables_found = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Check for tables
            for table in doc.tables:
                tables_found += 1
            
            return {
                'type': 'docx',
                'paragraphs': len(full_text),
                'tables': tables_found,
                'sample_text': '\n'.join(full_text[:10]),
                'file_size_kb': round(os.path.getsize(file_path) / 1024, 2)
            }
        except Exception as e:
            return {'type': 'docx', 'error': str(e)}


class TextDataSourceAdapter(DataSourceAdapter):
    """Text file data source adapter"""
    
    def connect(self) -> bool:
        """Check if text file exists and is readable"""
        try:
            file_path = self.source_config.get('file_path')
            if not file_path:
                return False
            
            with open(file_path, 'r') as f:
                f.readline()
            return True
        except Exception as e:
            logger.error(f"Text connection failed: {str(e)}")
            return False
    
    def extract_data(self, limit: Optional[int] = None) -> List[Dict[str, Any]]:
        """Extract data from text file"""
        try:
            file_path = self.source_config.get('file_path')
            encoding = self.source_config.get('encoding', 'utf-8')
            structure = self.source_config.get('structure', {})
            
            if structure.get('type') == 'pipe_delimited':
                delimiter = '|'
            elif structure.get('type') == 'tab_delimited':
                delimiter = '\t'
            elif structure.get('type') == 'comma_delimited':
                delimiter = ','
            else:
                # Try to detect delimiter from first line
                with open(file_path, 'r', encoding=encoding) as f:
                    first_line = f.readline().strip()
                    if '|' in first_line:
                        delimiter = '|'
                    elif '\t' in first_line:
                        delimiter = '\t'
                    elif ',' in first_line:
                        delimiter = ','
                    else:
                        return []
            
            data = []
            with open(file_path, 'r', encoding=encoding) as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                for i, row in enumerate(reader):
                    if limit and i >= limit:
                        break
                    # Clean and prepare data
                    cleaned_row = {k.strip(): v.strip() if isinstance(v, str) else v 
                                 for k, v in row.items()}
                    data.append(cleaned_row)
            
            return data
            
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            return []
    
    def get_schema(self) -> Dict[str, Any]:
        """Get text file schema information"""
        try:
            file_path = self.source_config.get('file_path')
            encoding = self.source_config.get('encoding', 'utf-8')
            
            with open(file_path, 'r', encoding=encoding) as f:
                first_line = f.readline().strip()
                
            # Check for common delimiters
            if '|' in first_line:
                columns = [col.strip() for col in first_line.split('|')]
                return {
                    'type': 'pipe_delimited',
                    'columns': columns,
                    'total_columns': len(columns),
                    'delimiter': '|'
                }
            elif '\t' in first_line:
                columns = [col.strip() for col in first_line.split('\t')]
                return {
                    'type': 'tab_delimited',
                    'columns': columns,
                    'total_columns': len(columns),
                    'delimiter': '\t'
                }
            elif ',' in first_line and first_line.count(',') > 2:
                columns = [col.strip() for col in first_line.split(',')]
                return {
                    'type': 'comma_delimited',
                    'columns': columns,
                    'total_columns': len(columns),
                    'delimiter': ','
                }
            else:
                return {
                    'type': 'plain_text',
                    'columns': [],
                    'total_columns': 0
                }
        except Exception as e:
            return {'type': 'text', 'error': str(e)}


# Global integration manager instance
integration_manager = DataSourceIntegrationManager()
