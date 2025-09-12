"""
Intelligent Data Scraper
Automatically discovers and processes all data files in the project
"""

import os
import re
import json
import csv
import logging
import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
from django.conf import settings
from django.db import transaction
from .models import DataSource, RawDataRecord, DataProcessingEvent
from .data_source_integration import integration_manager
from .data_validation import DataValidator
from .ai_geolocation import AIGeolocationEnhancer

logger = logging.getLogger(__name__)


class IntelligentDataScraper:
    """Intelligent data scraper that discovers and processes all data files"""
    
    def __init__(self):
        self.validator = DataValidator()
        self.geolocation_enhancer = AIGeolocationEnhancer()
        self.processed_files = set()
        self.data_directories = [
            'facilities_import/data/raw',
            'facilities_import/data',
            'facilities_import/src/data',
            'facilities_import/backup/backup_20250903_112502/src/data'
        ]
        self.supported_extensions = {
            '.csv': 'csv',
            '.json': 'json',
            '.xlsx': 'excel',
            '.xls': 'excel',
            '.xml': 'xml',
            '.txt': 'text',
            '.tsv': 'tsv',
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc'
        }
    
    def discover_data_files(self) -> List[Dict[str, Any]]:
        """Discover all data files in the project"""
        discovered_files = []
        
        for directory in self.data_directories:
            if os.path.exists(directory):
                for root, dirs, files in os.walk(directory):
                    for file in files:
                        file_path = os.path.join(root, file)
                        file_ext = os.path.splitext(file)[1].lower()
                        
                        if file_ext in self.supported_extensions:
                            file_info = self._analyze_file(file_path)
                            if file_info:
                                discovered_files.append(file_info)
        
        # Also check for files in the project root
        for file in os.listdir('.'):
            if os.path.isfile(file):
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in self.supported_extensions:
                    file_info = self._analyze_file(file)
                    if file_info:
                        discovered_files.append(file_info)
        
        logger.info(f"Discovered {len(discovered_files)} data files")
        return discovered_files
    
    def _analyze_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Analyze a file to determine its structure and content"""
        try:
            file_size = os.path.getsize(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            file_name = os.path.basename(file_path)
            
            # Skip if file is too small or already processed
            if file_size < 100 or file_path in self.processed_files:
                return None
            
            file_info = {
                'path': file_path,
                'name': file_name,
                'size': file_size,
                'extension': file_ext,
                'type': self.supported_extensions.get(file_ext, 'unknown'),
                'modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
                'encoding': self._detect_encoding(file_path),
                'structure': self._analyze_structure(file_path, file_ext)
            }
            
            return file_info
            
        except Exception as e:
            logger.warning(f"Failed to analyze file {file_path}: {str(e)}")
            return None
    
    def _detect_encoding(self, file_path: str) -> str:
        """Detect file encoding"""
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8')
        except:
            return 'utf-8'
    
    def _analyze_structure(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        """Analyze file structure to understand data format"""
        try:
            if file_ext == '.csv':
                return self._analyze_csv_structure(file_path)
            elif file_ext == '.json':
                return self._analyze_json_structure(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._analyze_excel_structure(file_path)
            elif file_ext == '.txt':
                return self._analyze_text_structure(file_path)
            elif file_ext == '.pdf':
                return self._analyze_pdf_structure(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._analyze_docx_structure(file_path)
            else:
                return {'type': 'unknown', 'columns': [], 'sample_rows': 0}
        except Exception as e:
            logger.warning(f"Failed to analyze structure of {file_path}: {str(e)}")
            return {'type': 'unknown', 'columns': [], 'sample_rows': 0, 'error': str(e)}
    
    def _analyze_csv_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze CSV file structure"""
        try:
            with open(file_path, 'r', encoding=self._detect_encoding(file_path)) as f:
                reader = csv.DictReader(f)
                columns = reader.fieldnames or []
                
                # Sample first few rows
                sample_rows = []
                for i, row in enumerate(reader):
                    if i >= 5:  # Sample first 5 rows
                        break
                    sample_rows.append(row)
                
                return {
                    'type': 'csv',
                    'columns': columns,
                    'total_columns': len(columns),
                    'sample_rows': len(sample_rows),
                    'delimiter': ','  # Could be enhanced to detect delimiter
                }
        except Exception as e:
            return {'type': 'csv', 'error': str(e)}
    
    def _analyze_json_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze JSON file structure"""
        try:
            with open(file_path, 'r', encoding=self._detect_encoding(file_path)) as f:
                data = json.load(f)
                
                if isinstance(data, list):
                    if data:
                        sample_record = data[0]
                        columns = list(sample_record.keys()) if isinstance(sample_record, dict) else []
                        return {
                            'type': 'json_array',
                            'columns': columns,
                            'total_columns': len(columns),
                            'total_records': len(data),
                            'sample_rows': min(5, len(data))
                        }
                    else:
                        return {'type': 'json_array', 'columns': [], 'total_records': 0}
                elif isinstance(data, dict):
                    columns = list(data.keys())
                    return {
                        'type': 'json_object',
                        'columns': columns,
                        'total_columns': len(columns),
                        'total_records': 1
                    }
                else:
                    return {'type': 'json_primitive', 'columns': [], 'total_records': 1}
        except Exception as e:
            return {'type': 'json', 'error': str(e)}
    
    def _analyze_excel_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze Excel file structure"""
        try:
            import pandas as pd
            df = pd.read_excel(file_path, nrows=5)  # Read first 5 rows
            return {
                'type': 'excel',
                'columns': df.columns.tolist(),
                'total_columns': len(df.columns),
                'sample_rows': len(df),
                'sheets': []  # Could be enhanced to list all sheets
            }
        except Exception as e:
            return {'type': 'excel', 'error': str(e)}
    
    def _analyze_text_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze text file structure to detect delimiters"""
        try:
            with open(file_path, 'r', encoding=self._detect_encoding(file_path)) as f:
                first_line = f.readline().strip()
                
            # Check for common delimiters
            if '|' in first_line:
                # Pipe-delimited
                columns = [col.strip() for col in first_line.split('|')]
                return {
                    'type': 'pipe_delimited',
                    'columns': columns,
                    'total_columns': len(columns),
                    'delimiter': '|'
                }
            elif '\t' in first_line:
                # Tab-delimited
                columns = [col.strip() for col in first_line.split('\t')]
                return {
                    'type': 'tab_delimited',
                    'columns': columns,
                    'total_columns': len(columns),
                    'delimiter': '\t'
                }
            elif ',' in first_line and first_line.count(',') > 2:
                # Comma-delimited
                columns = [col.strip() for col in first_line.split(',')]
                return {
                    'type': 'comma_delimited',
                    'columns': columns,
                    'total_columns': len(columns),
                    'delimiter': ','
                }
            else:
                # Plain text
                return {
                    'type': 'plain_text',
                    'columns': [],
                    'total_columns': 0
                }
        except Exception as e:
            return {'type': 'text', 'error': str(e)}
    
    def _analyze_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze PDF file structure"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Extract text from first page to analyze structure
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text()
                
                # Look for common patterns
                lines = text.split('\n')
                potential_headers = []
                for line in lines[:10]:  # Check first 10 lines
                    if line.strip() and len(line.split()) > 2:
                        potential_headers.append(line.strip())
                
                return {
                    'type': 'pdf',
                    'pages': num_pages,
                    'sample_text': text[:500],  # First 500 characters
                    'potential_headers': potential_headers[:5],
                    'file_size_mb': round(os.path.getsize(file_path) / (1024 * 1024), 2)
                }
        except Exception as e:
            return {'type': 'pdf', 'error': str(e)}
    
    def _analyze_docx_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze DOCX file structure"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text and analyze structure
            full_text = []
            tables_found = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Check for tables
            for table in doc.tables:
                tables_found += 1
                # Extract table headers if available
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    if headers:
                        full_text.append(f"Table {tables_found} headers: {', '.join(headers)}")
            
            return {
                'type': 'docx',
                'paragraphs': len(full_text),
                'tables': tables_found,
                'sample_text': '\n'.join(full_text[:10]),  # First 10 paragraphs
                'file_size_kb': round(os.path.getsize(file_path) / 1024, 2)
            }
        except Exception as e:
            return {'type': 'docx', 'error': str(e)}
    
    def create_data_sources(self, discovered_files: List[Dict[str, Any]]) -> List[DataSource]:
        """Create data sources for discovered files"""
        created_sources = []
        
        for file_info in discovered_files:
            try:
                # Generate source name from file path
                source_name = self._generate_source_name(file_info['path'])
                
                # Create configuration based on file type
                config = self._create_source_config(file_info)
                
                # Create data source
                data_source = DataSource.objects.create(
                    name=source_name,
                    source_type=file_info['type'],
                    description=f"Auto-discovered from {file_info['path']}",
                    configuration=config,
                    is_active=True
                )
                
                created_sources.append(data_source)
                logger.info(f"Created data source: {source_name}")
                
            except Exception as e:
                logger.error(f"Failed to create data source for {file_info['path']}: {str(e)}")
        
        return created_sources
    
    def _generate_source_name(self, file_path: str) -> str:
        """Generate a meaningful name for the data source"""
        file_name = os.path.basename(file_path)
        name_without_ext = os.path.splitext(file_name)[0]
        
        # Clean up the name
        name = name_without_ext.replace('_', ' ').replace('-', ' ').title()
        
        # Add directory context if needed
        if 'facilities' in file_path.lower():
            name = f"Facilities - {name}"
        elif 'counties' in file_path.lower():
            name = f"Counties - {name}"
        elif 'exports' in file_path.lower():
            name = f"Export - {name}"
        
        # Add timestamp to ensure uniqueness
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        name = f"{name} - {timestamp}"
        
        return name
    
    def _create_source_config(self, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Create configuration for data source based on file info"""
        config = {
            'file_path': file_info['path'],
            'encoding': file_info['encoding'],
            'file_size': file_info['size'],
            'discovered_at': datetime.now().isoformat()
        }
        
        if file_info['type'] == 'csv':
            config.update({
                'delimiter': ',',
                'has_header': True
            })
        elif file_info['type'] == 'json':
            config.update({
                'array_key': None,  # Will be detected during processing
                'structure_type': file_info['structure'].get('type', 'unknown')
            })
        elif file_info['type'] in ['excel', 'xlsx', 'xls']:
            config.update({
                'sheet_name': 0,  # First sheet
                'has_header': True
            })
        
        return config
    
    def process_all_files(self) -> Dict[str, Any]:
        """Process all discovered data files"""
        try:
            # Discover files
            discovered_files = self.discover_data_files()
            if not discovered_files:
                return {
                    'status': 'warning',
                    'message': 'No data files found to process',
                    'files_processed': 0
                }
            
            # Create data sources
            data_sources = self.create_data_sources(discovered_files)
            
            # Process each file
            results = {
                'status': 'success',
                'files_discovered': len(discovered_files),
                'sources_created': len(data_sources),
                'files_processed': 0,
                'records_ingested': 0,
                'errors': []
            }
            
            for data_source in data_sources:
                try:
                    # Process the file
                    file_result = self.process_file(data_source)
                    results['files_processed'] += 1
                    results['records_ingested'] += file_result.get('records', 0)
                    
                    # Mark file as processed
                    self.processed_files.add(data_source.configuration['file_path'])
                    
                except Exception as e:
                    error_msg = f"Failed to process {data_source.name}: {str(e)}"
                    results['errors'].append(error_msg)
                    logger.error(error_msg)
            
            logger.info(f"Processed {results['files_processed']} files, ingested {results['records_ingested']} records")
            return results
            
        except Exception as e:
            logger.error(f"Data scraping failed: {str(e)}")
            return {
                'status': 'error',
                'message': str(e),
                'files_processed': 0
            }
    
    def process_file(self, data_source: DataSource) -> Dict[str, Any]:
        """Process a single data file"""
        try:
            file_path = data_source.configuration['file_path']
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Extract data based on file type
            if file_ext == '.csv':
                data = self._extract_csv_data(file_path, data_source.configuration)
            elif file_ext == '.json':
                data = self._extract_json_data(file_path, data_source.configuration)
            elif file_ext in ['.xlsx', '.xls']:
                data = self._extract_excel_data(file_path, data_source.configuration)
            elif file_ext == '.txt':
                data = self._extract_text_data(file_path, data_source.configuration)
            elif file_ext == '.pdf':
                data = self._extract_pdf_data(file_path, data_source.configuration)
            elif file_ext in ['.docx', '.doc']:
                data = self._extract_docx_data(file_path, data_source.configuration)
            else:
                return {'status': 'skipped', 'message': f'Unsupported file type: {file_ext}'}
            
            if not data:
                return {'status': 'warning', 'message': 'No data extracted', 'records': 0}
            
            # Validate data
            validation_results = self.validator.validate_batch(data)
            
            # Enhance with geolocation (temporarily disabled for testing)
            # enhanced_data = self.geolocation_enhancer.enhance_batch(data)
            enhanced_data = data  # Skip geolocation enhancement for now
            
            # Store in raw data lake
            stored_records = []
            for record in enhanced_data:
                # Generate unique data_id and checksum
                data_id = f"{data_source.id}_{len(stored_records)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                checksum = hashlib.sha256(json.dumps(record, sort_keys=True).encode()).hexdigest()
                
                raw_record = RawDataRecord.objects.create(
                    source=data_source,
                    data_id=data_id,
                    raw_data=record,
                    metadata={
                        'file_path': file_path,
                        'validation_results': validation_results,
                        'processing_time': datetime.now().isoformat()
                    },
                    checksum=checksum,
                    processing_status='completed'
                )
                stored_records.append(raw_record.id)
            
            # Update data source last sync
            data_source.last_sync = datetime.now()
            data_source.save()
            
            # Log processing event
            DataProcessingEvent.objects.create(
                event_type='file_processing',
                record_id=f"file_{data_source.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                source=data_source,
                event_data={
                    'file_path': file_path,
                    'records_processed': len(stored_records),
                    'validation_results': validation_results
                },
                success=True
            )
            
            return {
                'status': 'success',
                'records': len(stored_records),
                'validation_results': validation_results
            }
            
        except Exception as e:
            logger.error(f"Failed to process file {data_source.name}: {str(e)}")
            
            # Log error event
            DataProcessingEvent.objects.create(
                event_type='file_processing',
                record_id=f"error_{data_source.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                source=data_source,
                event_data={'error': str(e)},
                success=False,
                error_message=str(e)
            )
            
            return {
                'status': 'error',
                'message': str(e),
                'records': 0
            }
    
    def _extract_csv_data(self, file_path: str, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract data from CSV file"""
        try:
            encoding = config.get('encoding', 'utf-8')
            delimiter = config.get('delimiter', ',')
            
            data = []
            with open(file_path, 'r', encoding=encoding) as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                for row in reader:
                    # Clean and prepare data
                    cleaned_row = {k.strip(): v.strip() if isinstance(v, str) else v 
                                 for k, v in row.items()}
                    data.append(cleaned_row)
            
            return data
        except Exception as e:
            logger.error(f"CSV extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_json_data(self, file_path: str, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract data from JSON file"""
        try:
            encoding = config.get('encoding', 'utf-8')
            
            with open(file_path, 'r', encoding=encoding) as f:
                data = json.load(f)
            
            # Handle different JSON structures
            if isinstance(data, list):
                return data
            elif isinstance(data, dict):
                # Look for common data keys
                for key in ['data', 'results', 'items', 'facilities', 'records']:
                    if key in data and isinstance(data[key], list):
                        return data[key]
                # If no array found, return the dict as single record
                return [data]
            else:
                return [{'value': data}]
                
        except Exception as e:
            logger.error(f"JSON extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_excel_data(self, file_path: str, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract data from Excel file with intelligent parsing for different file types"""
        try:
            import pandas as pd
            
            # Check file type based on filename
            filename = os.path.basename(file_path).lower()
            
            if 'fgm' in filename:
                return self._extract_fgm_resources(file_path)
            elif 'gbv' in filename and 'station' in filename:
                return self._extract_gbv_stations(file_path)
            elif 'police' in filename:
                return self._extract_police_stations(file_path)
            else:
                # Generic Excel extraction
                sheet_name = config.get('sheet_name', 0)
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to list of dictionaries
                data = df.to_dict('records')
                
                # Clean data - set None for missing values instead of empty strings
                cleaned_data = []
                for record in data:
                    cleaned_record = {}
                    for k, v in record.items():
                        if pd.isna(v):
                            cleaned_record[k] = None
                        elif isinstance(v, (int, float)) and pd.isna(v):
                            cleaned_record[k] = None
                        else:
                            cleaned_record[k] = v
                    cleaned_record['source_file'] = os.path.basename(file_path)
                    cleaned_record['source_type'] = 'excel_extraction'
                    cleaned_data.append(cleaned_record)
                
                return cleaned_data
            
        except Exception as e:
            logger.error(f"Excel extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_text_data(self, file_path: str, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract data from text file"""
        try:
            encoding = config.get('encoding', 'utf-8')
            structure = config.get('structure', {})
            
            if structure.get('type') == 'pipe_delimited':
                delimiter = '|'
            elif structure.get('type') == 'tab_delimited':
                delimiter = '\t'
            elif structure.get('type') == 'comma_delimited':
                delimiter = ','
            else:
                # Try to detect delimiter from first line
                with open(file_path, 'r', encoding=encoding) as f:
                    first_line = f.readline().strip()
                    if '|' in first_line:
                        delimiter = '|'
                    elif '\t' in first_line:
                        delimiter = '\t'
                    elif ',' in first_line:
                        delimiter = ','
                    else:
                        return []
            
            data = []
            with open(file_path, 'r', encoding=encoding) as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                for row in reader:
                    # Clean and prepare data
                    cleaned_row = {k.strip(): v.strip() if isinstance(v, str) else v 
                                 for k, v in row.items()}
                    data.append(cleaned_row)
            
            return data
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_pdf_data(self, file_path: str, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract data from PDF file with intelligent parsing for different PDF types"""
        try:
            import PyPDF2
            import re
            
            data = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Process all pages for comprehensive data extraction
                max_pages = len(pdf_reader.pages)
                
                all_text = ""
                for page_num in range(max_pages):
                    page = pdf_reader.pages[page_num]
                    all_text += page.extract_text() + "\n"
                
                # Check if this is KMPDC facilities data
                if "LIST OF LICENCED PUBLIC HEALTH FACILITIES" in all_text:
                    facilities = self._extract_facilities_from_pdf_text(all_text)
                    for facility in facilities:
                        facility['source_file'] = os.path.basename(file_path)
                        facility['source_type'] = 'pdf_extraction'
                    data.extend(facilities)
                elif "NATIONAL SHELTERS NETWORK" in all_text:
                    shelters = self._extract_shelters_from_pdf_text(all_text)
                    for shelter in shelters:
                        shelter['source_file'] = os.path.basename(file_path)
                        shelter['source_type'] = 'pdf_extraction'
                    data.extend(shelters)
                else:
                    # General text extraction
                    lines = all_text.split('\n')
                    for line in lines:
                        if line.strip() and len(line.split()) > 3:
                            data.append({
                                'text': line.strip(),
                                'source_file': os.path.basename(file_path),
                                'source_type': 'pdf_extraction'
                            })
            
            return data
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_docx_data(self, file_path: str, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract data from DOCX file with intelligent parsing for GBV organizations"""
        try:
            from docx import Document
            
            data = []
            doc = Document(file_path)
            
            # Check if this is a GBV organizations file
            is_gbv_file = 'gbv' in file_path.lower() or 'support' in file_path.lower()
            
            if is_gbv_file:
                # Parse GBV organizations with structured data
                data = self._extract_gbv_organizations(doc, file_path)
            else:
                # Standard DOCX extraction
                # Extract text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        data.append({
                            'text': para.text.strip(),
                            'type': 'paragraph',
                            'source_file': os.path.basename(file_path)
                        })
                
                # Extract data from tables
                for table_num, table in enumerate(doc.tables):
                    if table.rows:
                        # Get headers from first row
                        headers = [cell.text.strip() for cell in table.rows[0].cells]
                        
                        # Extract data from remaining rows
                        for row in table.rows[1:]:
                            row_data = {}
                            for i, cell in enumerate(row.cells):
                                if i < len(headers):
                                    row_data[headers[i]] = cell.text.strip()
                            if row_data:
                                row_data['table_number'] = table_num + 1
                                row_data['source_file'] = os.path.basename(file_path)
                                data.append(row_data)
            
            return data
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_gbv_organizations(self, doc, file_path: str) -> List[Dict[str, Any]]:
        """Extract GBV organizations with structured data mapping"""
        organizations = []
        current_org = {}
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a new organization (numbered list)
            if re.match(r'^\d+\.\s+[A-Z]', text):
                # Save previous organization if exists
                if current_org:
                    organizations.append(current_org)
                
                # Start new organization
                org_name = re.sub(r'^\d+\.\s+', '', text)
                current_org = {
                    'organization_name': org_name,
                    'source_file': os.path.basename(file_path),
                    'source_type': 'docx_extraction',
                    'services': [],
                    'coverage': '',
                    'contact_details': {},
                    'main_offices': [],
                    'toll_free_number': '',
                    'website': '',
                    'experience': '',
                    'beneficiaries': '',
                    'support_network': []
                }
                current_section = None
            
            # Check for section headers
            elif text.startswith('**Services**'):
                current_section = 'services'
            elif text.startswith('**Coverage**'):
                current_section = 'coverage'
            elif text.startswith('**Contact Details**'):
                current_section = 'contact'
            elif text.startswith('**Main Offices**') or text.startswith('**Offices**'):
                current_section = 'offices'
            elif text.startswith('**Toll-Free Number**'):
                current_section = 'toll_free'
            elif text.startswith('**Website**'):
                current_section = 'website'
            elif text.startswith('**Experience**'):
                current_section = 'experience'
            elif text.startswith('**Beneficiaries**'):
                current_section = 'beneficiaries'
            elif text.startswith('**Support Network**'):
                current_section = 'support_network'
            
            # Process content based on current section
            elif current_org and current_section:
                if current_section == 'services':
                    # Extract services (comma-separated)
                    services = [s.strip() for s in text.split(',')]
                    current_org['services'].extend(services)
                elif current_section == 'coverage':
                    current_org['coverage'] = text
                elif current_section == 'contact':
                    # Extract contact details
                    if ':' in text:
                        key, value = text.split(':', 1)
                        current_org['contact_details'][key.strip()] = value.strip()
                elif current_section == 'offices':
                    # Extract office locations
                    if text.startswith('*'):
                        office = text.strip('*').strip()
                        current_org['main_offices'].append(office)
                elif current_section == 'toll_free':
                    current_org['toll_free_number'] = text
                elif current_section == 'website':
                    current_org['website'] = text
                elif current_section == 'experience':
                    current_org['experience'] = text
                elif current_section == 'beneficiaries':
                    current_org['beneficiaries'] = text
                elif current_section == 'support_network':
                    # Extract support network items
                    if text.startswith('*'):
                        network_item = text.strip('*').strip()
                        current_org['support_network'].append(network_item)
        
        # Add the last organization
        if current_org:
            organizations.append(current_org)
        
        return organizations
    
    def _extract_facilities_from_pdf_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract facility information from KMPDC PDF text using regex patterns"""
        facilities = []
        
        # Split text into lines
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Skip header lines
            if 'LIST OF LICENCED' in line or 'S/NO Reg_No' in line:
                continue
            
            # Look for lines that contain GK- registration numbers
            # Pattern: numberGK-XXXXXX or GK-XXXXXX followed by facility data
            if re.search(r'GK-\d+', line) and ('LICENCED' in line or 'LEVEL' in line):
                try:
                    # Use a more sophisticated parsing approach
                    # The structure is: S_NO Reg_No Facility_Name Facility_Type Facility_Agent Level County Status
                    
                    # Extract sequence number and registration number
                    # Handle both patterns: "100 GK-012180" and "GK-012180"
                    s_no_match = re.match(r'^(\d+)\s+(GK-\d+)', line)
                    if s_no_match:
                        s_no = s_no_match.group(1)
                        reg_no = s_no_match.group(2)
                        remaining_line = line[len(s_no + ' ' + reg_no):].strip()
                    else:
                        # Try pattern without sequence number
                        gk_match = re.search(r'(GK-\d+)', line)
                        if not gk_match:
                            continue
                        s_no = "0"  # Default sequence number
                        reg_no = gk_match.group(1)
                        remaining_line = line.replace(reg_no, '', 1).strip()
                    
                    # Find the last two words (County and Status)
                    words = remaining_line.split()
                    if len(words) < 4:
                        continue
                    
                    status = words[-1]  # Last word is status
                    county = words[-2]  # Second to last word is county
                    
                    # Remove county and status from the line
                    facility_data = ' '.join(words[:-2])
                    
                    # Now we need to parse: Facility_Name Facility_Type Facility_Agent Level
                    # Look for the LAST LEVEL pattern (there might be duplicates)
                    level_matches = list(re.finditer(r'(LEVEL\s+\d+[A-Z]?)', facility_data))
                    if not level_matches:
                        continue
                    
                    # Use the last LEVEL match
                    level_match = level_matches[-1]
                    level = level_match.group(1)
                    facility_info = facility_data[:level_match.start()].strip()
                    
                    # Now parse facility_info: Facility_Name Facility_Type Facility_Agent
                    # Look for facility types
                    facility_types = ['HOSPITAL', 'DISPENSARY', 'HEALTH CENTRE', 'CLINIC', 'MEDICAL CENTER', 'BASIC HEALTH CENTRE', 'MEDICAL CLINIC']
                    
                    facility_type = None
                    facility_type_index = -1
                    facility_info_words = facility_info.split()
                    
                    for i, word in enumerate(facility_info_words):
                        if any(ft in word.upper() for ft in facility_types):
                            facility_type = word
                            facility_type_index = i
                            break
                    
                    if facility_type_index == -1:
                        continue
                    
                    # Extract facility name (everything before facility type)
                    facility_name = ' '.join(facility_info_words[:facility_type_index])
                    
                    # Extract facility agent (everything after facility type)
                    facility_agent = ' '.join(facility_info_words[facility_type_index + 1:])
                    
                    if facility_name and facility_type and county and status:
                        facility = {
                            's_no': s_no,
                            'registration_number': reg_no,
                            'facility_name': facility_name.strip(),
                            'facility_type': facility_type.strip(),
                            'facility_agent': facility_agent.strip() or 'Unknown',
                            'level': level.strip(),
                            'county': county.strip(),
                            'status': status.strip(),
                            'source_type': 'pdf_extraction'
                        }
                        facilities.append(facility)
                        
                except Exception as e:
                    logger.warning(f"Error parsing facility line: {line} - {str(e)}")
                    continue
        
        return facilities
    
    def _extract_shelters_from_pdf_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract shelter data from National Shelters Network PDF"""
        shelters = []
        lines = text.split('\n')
        
        # Find the start of data (after header)
        data_start = 0
        for i, line in enumerate(lines):
            if "NO." in line and "NAME OF SHELTER" in line and "COUNTY" in line:
                data_start = i + 1
                break
        
        # Process each line
        for line in lines[data_start:]:
            line = line.strip()
            if not line or len(line) < 10:  # Skip empty or very short lines
                continue
            
            # Skip header lines
            if "NO." in line or "NAME OF SHELTER" in line or "COUNTY" in line:
                continue
            
            # Extract shelter data
            shelter = self._parse_shelter_line(line)
            if shelter:
                shelters.append(shelter)
        
        return shelters
    
    def _parse_shelter_line(self, line: str) -> Optional[Dict[str, Any]]:
        """Parse a single shelter line from National Shelters Network PDF"""
        try:
            # Split by multiple spaces to separate fields
            parts = re.split(r'\s{2,}', line.strip())
            
            if len(parts) < 4:
                return None
            
            # Extract fields based on position
            no = parts[0].strip() if parts[0] else None
            shelter_name = parts[1].strip() if len(parts) > 1 else None
            county = parts[2].strip() if len(parts) > 2 else None
            operator = parts[3].strip() if len(parts) > 3 else None
            
            # Contact information is usually in the last parts
            contact = None
            if len(parts) > 4:
                contact = ' '.join(parts[4:]).strip()
            
            return {
                'shelter_number': no,
                'shelter_name': shelter_name,
                'county': county,
                'shelter_operator': operator,
                'contact_info': contact
            }
            
        except Exception as e:
            logger.error(f"Error parsing shelter line '{line}': {str(e)}")
            return None
    
    def _extract_fgm_resources(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract FGM resources data from Excel"""
        try:
            import pandas as pd
            data = []
            df = pd.read_excel(file_path, sheet_name=0)
            
            for _, row in df.iterrows():
                record = {
                    'document_title': row.get('Document title', None),
                    'document_type': row.get('Document type', None),
                    'description': row.get('Description ', None),
                    'file_url': row.get('File url ', None),
                    'file_name': row.get('File name', None),
                    'gbv_category': row.get('Gbv category ', None),
                    'image_url': row.get('image url', None),
                    'external_url': row.get('External url ', None),
                    'source_file': os.path.basename(file_path),
                    'source_type': 'excel_extraction'
                }
                
                # Remove None values to avoid misinformation
                record = {k: v for k, v in record.items() if v is not None and str(v).strip() != ''}
                data.append(record)
            
            return data
        except Exception as e:
            logger.error(f"FGM resources extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_gbv_stations(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract GBV station data from Excel"""
        try:
            import pandas as pd
            data = []
            df = pd.read_excel(file_path, sheet_name=0)
            
            for _, row in df.iterrows():
                record = {
                    'mapping_of_police_commands': row.get('MAPPING OF POLICE COMMANDS', None),
                    'kps_station_data': row.get('KPS STATION DATA', None),
                    'county': row.get('Unnamed: 2', None),
                    'sub_county': row.get('Unnamed: 3', None),
                    'police_station': row.get('Unnamed: 4', None),
                    'source_file': os.path.basename(file_path),
                    'source_type': 'excel_extraction'
                }
                
                # Remove None values to avoid misinformation
                record = {k: v for k, v in record.items() if v is not None and str(v).strip() != ''}
                data.append(record)
            
            return data
        except Exception as e:
            logger.error(f"GBV stations extraction failed for {file_path}: {str(e)}")
            return []
    
    def _extract_police_stations(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract police station data from Excel"""
        try:
            import pandas as pd
            data = []
            df = pd.read_excel(file_path, sheet_name=0)
            
            for _, row in df.iterrows():
                record = {
                    'police_station_name': row.get('NAIROBI', None),
                    'phone_number': row.get('Unnamed: 1', None),
                    'sub_county': row.get('Unnamed: 2', None),
                    'source_file': os.path.basename(file_path),
                    'source_type': 'excel_extraction'
                }
                
                # Remove None values to avoid misinformation
                record = {k: v for k, v in record.items() if v is not None and str(v).strip() != ''}
                data.append(record)
            
            return data
        except Exception as e:
            logger.error(f"Police stations extraction failed for {file_path}: {str(e)}")
            return []
    
    def clear_sample_data(self) -> Dict[str, Any]:
        """Clear all sample data from the system"""
        try:
            with transaction.atomic():
                # Clear sample data sources
                sample_sources = DataSource.objects.filter(
                    name__icontains='sample'
                ).delete()
                
                # Clear sample raw data records
                sample_records = RawDataRecord.objects.filter(
                    metadata__contains={'source_type': 'sample'}
                ).delete()
                
                # Clear sample processing events
                sample_events = DataProcessingEvent.objects.filter(
                    event_data__contains={'source_type': 'sample'}
                ).delete()
                
                logger.info(f"Cleared sample data: {sample_sources[0]} sources, {sample_records[0]} records, {sample_events[0]} events")
                
                return {
                    'status': 'success',
                    'message': 'Sample data cleared successfully',
                    'sources_deleted': sample_sources[0],
                    'records_deleted': sample_records[0],
                    'events_deleted': sample_events[0]
                }
                
        except Exception as e:
            logger.error(f"Failed to clear sample data: {str(e)}")
            return {
                'status': 'error',
                'message': str(e)
            }


# Global scraper instance
intelligent_scraper = IntelligentDataScraper()
