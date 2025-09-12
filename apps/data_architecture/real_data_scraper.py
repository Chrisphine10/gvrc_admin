"""
Real Data Scraper
Handles all actual data files from the raw directory with high accuracy
"""

import os
import re
import json
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import pandas as pd

logger = logging.getLogger(__name__)


class RealDataScraper:
    """Real data scraper for actual files in the raw directory"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.xls', '.csv', '.json', '.txt']
        self.raw_data_path = 'facilities_import/data/raw'
    
    def scrape_file(self, file_path: str) -> List[Dict[str, Any]]:
        """Main method to scrape data from any supported file"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return []
        
        if file_path.suffix.lower() not in self.supported_formats:
            logger.error(f"Unsupported file format: {file_path.suffix}")
            return []
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_pdf_data(str(file_path))
            elif file_path.suffix.lower() in ['.docx']:
                return self._extract_docx_data(str(file_path))
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                return self._extract_excel_data(str(file_path))
            elif file_path.suffix.lower() == '.csv':
                return self._extract_csv_data(str(file_path))
            elif file_path.suffix.lower() == '.json':
                return self._extract_json_data(str(file_path))
            elif file_path.suffix.lower() == '.txt':
                return self._extract_txt_data(str(file_path))
            else:
                logger.error(f"Unsupported file format: {file_path.suffix}")
                return []
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return []
    
    def _extract_pdf_data(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract data from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                # Process all pages for comprehensive data extraction
                max_pages = len(pdf_reader.pages)
                
                all_text = ""
                for page_num in range(max_pages):
                    page = pdf_reader.pages[page_num]
                    all_text += page.extract_text() + "\n"
                
                # Check if this is KMPDC facilities data
                if "LIST OF LICENCED PUBLIC HEALTH FACILITIES" in all_text:
                    return self._extract_facilities_from_pdf_text(all_text, file_path)
                elif "NATIONAL SHELTERS NETWORK" in all_text:
                    return self._extract_shelters_from_pdf_text(all_text, file_path)
                else:
                    # Generic PDF extraction
                    return self._extract_generic_pdf_data(all_text, file_path)
                    
        except Exception as e:
            logger.error(f"Error extracting PDF data from {file_path}: {str(e)}")
            return []
    
    def _extract_facilities_from_pdf_text(self, text: str, file_path: str) -> List[Dict[str, Any]]:
        """Extract facilities data from KMPDC PDF text"""
        facilities = []
        lines = text.split('\n')
        
        # Find the start of data (after header)
        data_start = 0
        for i, line in enumerate(lines):
            if "S/NO" in line and "Reg_No" in line and "Facility_Name" in line:
                data_start = i + 1
                break
        
        # Process each line
        for line in lines[data_start:]:
            line = line.strip()
            if not line or len(line) < 10:  # Skip empty or very short lines
                continue
            
            # Skip header lines
            if "S/NO" in line or "Reg_No" in line or "Facility_Name" in line:
                continue
            
            # Extract facility data using improved parsing
            facility = self._parse_facility_line(line, file_path)
            if facility:
                facilities.append(facility)
        
        return facilities
    
    def _parse_facility_line(self, line: str, file_path: str) -> Optional[Dict[str, Any]]:
        """Parse a single facility line from KMPDC PDF"""
        try:
            # Split by multiple spaces to separate fields
            parts = re.split(r'\s{2,}', line.strip())
            
            if len(parts) < 6:
                return None
            
            # Extract fields based on position and content
            s_no = parts[0].strip() if parts[0] else None
            reg_no = parts[1].strip() if len(parts) > 1 else None
            facility_name = parts[2].strip() if len(parts) > 2 else None
            facility_type = parts[3].strip() if len(parts) > 3 else None
            facility_agent = parts[4].strip() if len(parts) > 4 else None
            
            # Find level and county - look for "LEVEL" keyword
            level = None
            county = None
            status = None
            
            # Search for LEVEL in the remaining parts
            for i, part in enumerate(parts[5:], 5):
                if "LEVEL" in part:
                    level = part.strip()
                    # County should be the next part
                    if i + 1 < len(parts):
                        county = parts[i + 1].strip()
                    # Status should be the last part
                    if i + 2 < len(parts):
                        status = parts[i + 2].strip()
                    break
            
            # If no LEVEL found, try to extract from the last parts
            if not level and len(parts) >= 6:
                # Last part is usually status
                status = parts[-1].strip()
                # Second to last is usually county
                if len(parts) >= 7:
                    county = parts[-2].strip()
                # Look for level in the middle parts
                for part in parts[5:-2]:
                    if "LEVEL" in part:
                        level = part.strip()
                        break
            
            return {
                's_no': s_no,
                'registration_number': reg_no,
                'facility_name': facility_name,
                'facility_type': facility_type,
                'facility_agent': facility_agent,
                'level': level,
                'county': county,
                'status': status,
                'source_file': os.path.basename(file_path),
                'source_type': 'pdf_extraction'
            }
            
        except Exception as e:
            logger.error(f"Error parsing facility line '{line}': {str(e)}")
            return None
    
    def _extract_shelters_from_pdf_text(self, text: str, file_path: str) -> List[Dict[str, Any]]:
        """Extract shelter data from National Shelters Network PDF"""
        shelters = []
        lines = text.split('\n')
        
        # Find the start of data (after header)
        data_start = 0
        for i, line in enumerate(lines):
            if "NO." in line and "NAME OF SHELTER" in line and "COUNTY" in line:
                data_start = i + 1
                break
        
        # Process each line
        for line in lines[data_start:]:
            line = line.strip()
            if not line or len(line) < 10:  # Skip empty or very short lines
                continue
            
            # Skip header lines
            if "NO." in line or "NAME OF SHELTER" in line or "COUNTY" in line:
                continue
            
            # Extract shelter data
            shelter = self._parse_shelter_line(line, file_path)
            if shelter:
                shelters.append(shelter)
        
        return shelters
    
    def _parse_shelter_line(self, line: str, file_path: str) -> Optional[Dict[str, Any]]:
        """Parse a single shelter line from National Shelters Network PDF"""
        try:
            # Split by multiple spaces to separate fields
            parts = re.split(r'\s{2,}', line.strip())
            
            if len(parts) < 4:
                return None
            
            # Extract fields based on position
            no = parts[0].strip() if parts[0] else None
            shelter_name = parts[1].strip() if len(parts) > 1 else None
            county = parts[2].strip() if len(parts) > 2 else None
            operator = parts[3].strip() if len(parts) > 3 else None
            
            # Contact information is usually in the last parts
            contact = None
            if len(parts) > 4:
                contact = ' '.join(parts[4:]).strip()
            
            return {
                'shelter_number': no,
                'shelter_name': shelter_name,
                'county': county,
                'shelter_operator': operator,
                'contact_info': contact,
                'source_file': os.path.basename(file_path),
                'source_type': 'pdf_extraction'
            }
            
        except Exception as e:
            logger.error(f"Error parsing shelter line '{line}': {str(e)}")
            return None
    
    def _extract_generic_pdf_data(self, text: str, file_path: str) -> List[Dict[str, Any]]:
        """Extract generic data from PDF text"""
        return [{
            'content': text,
            'source_file': os.path.basename(file_path),
            'source_type': 'pdf_extraction'
        }]
    
    def _extract_docx_data(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract data from DOCX files"""
        try:
            doc = Document(file_path)
            
            # Check if this is GBV organizations document
            if "GBV Support Organizations" in doc.paragraphs[0].text:
                return self._extract_gbv_organizations(doc, file_path)
            else:
                # Generic DOCX extraction
                return self._extract_generic_docx_data(doc, file_path)
                
        except Exception as e:
            logger.error(f"Error extracting DOCX data from {file_path}: {str(e)}")
            return []
    
    def _extract_gbv_organizations(self, doc, file_path: str) -> List[Dict[str, Any]]:
        """Extract GBV organizations with structured data mapping"""
        organizations = []
        current_org = {}
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a new organization (numbered list)
            if re.match(r'^\d+\.\s+[A-Z]', text):
                # Save previous organization if exists
                if current_org:
                    organizations.append(current_org)
                
                # Start new organization
                org_name = re.sub(r'^\d+\.\s+', '', text)
                current_org = {
                    'organization_name': org_name,
                    'source_file': os.path.basename(file_path),
                    'source_type': 'docx_extraction',
                    'services': [],
                    'coverage': '',
                    'contact_details': {},
                    'main_offices': [],
                    'toll_free_number': '',
                    'website': '',
                    'experience': '',
                    'beneficiaries': '',
                    'support_network': []
                }
                current_section = None
            
            # Check for section headers
            elif text.startswith('**Services**'):
                current_section = 'services'
            elif text.startswith('**Coverage**'):
                current_section = 'coverage'
            elif text.startswith('**Contact Details**'):
                current_section = 'contact'
            elif text.startswith('**Main Offices**') or text.startswith('**Offices**'):
                current_section = 'offices'
            elif text.startswith('**Toll-Free Number**'):
                current_section = 'toll_free'
            elif text.startswith('**Website**'):
                current_section = 'website'
            elif text.startswith('**Experience**'):
                current_section = 'experience'
            elif text.startswith('**Beneficiaries**'):
                current_section = 'beneficiaries'
            elif text.startswith('**Support Network**'):
                current_section = 'support_network'
            
            # Process content based on current section
            elif current_org and current_section:
                if current_section == 'services':
                    # Extract services (comma-separated)
                    services = [s.strip() for s in text.split(',')]
                    current_org['services'].extend(services)
                elif current_section == 'coverage':
                    current_org['coverage'] = text
                elif current_section == 'contact':
                    # Extract contact details
                    if ':' in text:
                        key, value = text.split(':', 1)
                        current_org['contact_details'][key.strip()] = value.strip()
                elif current_section == 'offices':
                    # Extract office locations
                    if text.startswith('*'):
                        office = text.strip('*').strip()
                        current_org['main_offices'].append(office)
                elif current_section == 'toll_free':
                    current_org['toll_free_number'] = text
                elif current_section == 'website':
                    current_org['website'] = text
                elif current_section == 'experience':
                    current_org['experience'] = text
                elif current_section == 'beneficiaries':
                    current_org['beneficiaries'] = text
                elif current_section == 'support_network':
                    # Extract support network items
                    if text.startswith('*'):
                        network_item = text.strip('*').strip()
                        current_org['support_network'].append(network_item)
        
        # Add the last organization
        if current_org:
            organizations.append(current_org)
        
        return organizations
    
    def _extract_generic_docx_data(self, doc, file_path: str) -> List[Dict[str, Any]]:
        """Extract generic data from DOCX document"""
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        return [{
            'content': '\n'.join(content),
            'source_file': os.path.basename(file_path),
            'source_type': 'docx_extraction'
        }]
    
    def _extract_excel_data(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract data from Excel files with specific handling for different file types"""
        try:
            # Read all sheets
            xl = pd.ExcelFile(file_path)
            all_data = []
            
            # Check file type based on filename
            filename = os.path.basename(file_path).lower()
            
            if 'fgm' in filename:
                return self._extract_fgm_resources(xl, file_path)
            elif 'gbv' in filename and 'station' in filename:
                return self._extract_gbv_stations(xl, file_path)
            elif 'police' in filename:
                return self._extract_police_stations(xl, file_path)
            else:
                # Generic Excel extraction
                return self._extract_generic_excel_data(xl, file_path)
            
        except Exception as e:
            logger.error(f"Error extracting Excel data from {file_path}: {str(e)}")
            return []
    
    def _extract_fgm_resources(self, xl, file_path: str) -> List[Dict[str, Any]]:
        """Extract FGM resources data from Excel"""
        data = []
        df = pd.read_excel(file_path, sheet_name=0)
        
        for _, row in df.iterrows():
            record = {
                'document_title': row.get('Document title', None),
                'document_type': row.get('Document type', None),
                'description': row.get('Description ', None),
                'file_url': row.get('File url ', None),
                'file_name': row.get('File name', None),
                'gbv_category': row.get('Gbv category ', None),
                'image_url': row.get('image url', None),
                'external_url': row.get('External url ', None),
                'source_file': os.path.basename(file_path),
                'source_type': 'excel_extraction'
            }
            
            # Remove None values to avoid misinformation
            record = {k: v for k, v in record.items() if v is not None and str(v).strip() != ''}
            data.append(record)
        
        return data
    
    def _extract_gbv_stations(self, xl, file_path: str) -> List[Dict[str, Any]]:
        """Extract GBV station data from Excel"""
        data = []
        df = pd.read_excel(file_path, sheet_name=0)
        
        for _, row in df.iterrows():
            record = {
                'mapping_of_police_commands': row.get('MAPPING OF POLICE COMMANDS', None),
                'kps_station_data': row.get('KPS STATION DATA', None),
                'county': row.get('Unnamed: 2', None),
                'sub_county': row.get('Unnamed: 3', None),
                'police_station': row.get('Unnamed: 4', None),
                'source_file': os.path.basename(file_path),
                'source_type': 'excel_extraction'
            }
            
            # Remove None values to avoid misinformation
            record = {k: v for k, v in record.items() if v is not None and str(v).strip() != ''}
            data.append(record)
        
        return data
    
    def _extract_police_stations(self, xl, file_path: str) -> List[Dict[str, Any]]:
        """Extract police station data from Excel"""
        data = []
        df = pd.read_excel(file_path, sheet_name=0)
        
        for _, row in df.iterrows():
            record = {
                'police_station_name': row.get('NAIROBI', None),
                'phone_number': row.get('Unnamed: 1', None),
                'sub_county': row.get('Unnamed: 2', None),
                'source_file': os.path.basename(file_path),
                'source_type': 'excel_extraction'
            }
            
            # Remove None values to avoid misinformation
            record = {k: v for k, v in record.items() if v is not None and str(v).strip() != ''}
            data.append(record)
        
        return data
    
    def _extract_generic_excel_data(self, xl, file_path: str) -> List[Dict[str, Any]]:
        """Extract generic data from Excel files"""
        all_data = []
        
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Convert DataFrame to list of dictionaries
            for _, row in df.iterrows():
                record = {
                    'sheet_name': sheet_name,
                    'source_file': os.path.basename(file_path),
                    'source_type': 'excel_extraction'
                }
                
                # Add all columns as fields
                for col in df.columns:
                    value = row[col]
                    # Handle NaN values - set to None instead of empty string
                    if pd.isna(value):
                        record[col] = None
                    else:
                        record[col] = str(value)
                
                all_data.append(record)
        
        return all_data
    
    def _extract_csv_data(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract data from CSV files"""
        try:
            df = pd.read_csv(file_path)
            data = []
            
            for _, row in df.iterrows():
                record = {
                    'source_file': os.path.basename(file_path),
                    'source_type': 'csv_extraction'
                }
                
                # Add all columns as fields
                for col in df.columns:
                    value = row[col]
                    # Handle NaN values - set to None instead of empty string
                    if pd.isna(value):
                        record[col] = None
                    else:
                        record[col] = str(value)
                
                data.append(record)
            
            return data
            
        except Exception as e:
            logger.error(f"Error extracting CSV data from {file_path}: {str(e)}")
            return []
    
    def _extract_json_data(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract data from JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Handle different JSON structures
            if isinstance(data, list):
                for item in data:
                    item['source_file'] = os.path.basename(file_path)
                    item['source_type'] = 'json_extraction'
                return data
            elif isinstance(data, dict):
                data['source_file'] = os.path.basename(file_path)
                data['source_type'] = 'json_extraction'
                return [data]
            else:
                return [{
                    'content': str(data),
                    'source_file': os.path.basename(file_path),
                    'source_type': 'json_extraction'
                }]
                
        except Exception as e:
            logger.error(f"Error extracting JSON data from {file_path}: {str(e)}")
            return []
    
    def _extract_txt_data(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract data from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return [{
                'content': content,
                'source_file': os.path.basename(file_path),
                'source_type': 'txt_extraction'
            }]
            
        except Exception as e:
            logger.error(f"Error extracting TXT data from {file_path}: {str(e)}")
            return []
    
    def discover_all_files(self) -> List[Dict[str, Any]]:
        """Discover all files in the raw data directory"""
        discovered_files = []
        
        if not os.path.exists(self.raw_data_path):
            logger.error(f"Raw data directory not found: {self.raw_data_path}")
            return discovered_files
        
        for file_name in os.listdir(self.raw_data_path):
            file_path = os.path.join(self.raw_data_path, file_name)
            
            if os.path.isfile(file_path):
                file_ext = os.path.splitext(file_name)[1].lower()
                
                if file_ext in self.supported_formats:
                    file_info = {
                        'path': file_path,
                        'name': file_name,
                        'size': os.path.getsize(file_path),
                        'extension': file_ext,
                        'type': self.supported_extensions.get(file_ext, 'unknown')
                    }
                    discovered_files.append(file_info)
        
        logger.info(f"Discovered {len(discovered_files)} files in raw data directory")
        return discovered_files


# Global scraper instance
real_data_scraper = RealDataScraper()
